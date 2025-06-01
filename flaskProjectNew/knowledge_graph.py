import json
from openai import OpenAI
import os
import PyPDF2
from docx import Document
import math
import time
import traceback

# 配置qwen模型api
client = OpenAI(
    base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
    api_key=os.environ.get('QWEN_API_KEY', "sk-5fb4c90206e84a7b9d88fc91e72fbe44")
)

def read_file_content(file_path):
    """读取文件内容，支持PDF和DOCX格式"""
    try:
        print(f"开始读取文件：{file_path}")
        if not os.path.exists(file_path):
            print(f"文件不存在：{file_path}")
            # 尝试替代路径
            alternate_paths = [
                file_path.replace('\\uploads\\', '\\flaskProject\\static\\uploads\\'),
                file_path.replace('\\uploads\\', '\\static\\uploads\\'),
                file_path.replace('\\static\\uploads\\', '\\uploads\\')
            ]
            for alt_path in alternate_paths:
                print(f"尝试替代路径: {alt_path}")
                if os.path.exists(alt_path):
                    print(f"找到文件: {alt_path}")
                    file_path = alt_path
                    break
            else:
                return None  # 如果所有替代路径都不存在，返回None
        
        # 检查文件是否为PDF文件，支持没有扩展名的情况
        if file_path.endswith('.pdf') or file_path.endswith('_pdf') or '_pdf' in file_path:
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    k = 0
                    content = ""
                    for page in reader.pages:
                        if k < 30:
                            page_text = page.extract_text()
                            if page_text:
                                content += page_text + '\n\n'
                        k += 1
                    print(f"成功读取PDF文件，共{k}页，提取了前30页")
                    return content
            except Exception as e:
                print(f"PDF文件读取错误：{str(e)}")
                print(traceback.format_exc())
                return None
        # 检查文件是否为DOCX文件，支持没有扩展名的情况
        elif file_path.endswith('.docx') or file_path.endswith('_docx') or '_docx' in file_path:
            try:
                doc_text = ""
                doc = Document(file_path)
                paragraphs = doc.paragraphs
                for paragraph in paragraphs:
                    doc_text += paragraph.text + "\n"
                print(f"成功读取DOCX文件，共{len(paragraphs)}段落")
                return doc_text
            except Exception as e:
                print(f"DOCX文件读取错误：{str(e)}")
                print(traceback.format_exc())
                return None
        else:
            # 尝试以二进制方式打开文件检测类型
            try:
                with open(file_path, 'rb') as file:
                    header = file.read(4)
                    # PDF文件的前4个字节是"%PDF"
                    if header.startswith(b'%PDF'):
                        print("检测到PDF文件格式，尝试读取...")
                        file.seek(0)
                        reader = PyPDF2.PdfReader(file)
                        k = 0
                        content = ""
                        for page in reader.pages:
                            if k < 30:
                                page_text = page.extract_text()
                                if page_text:
                                    content += page_text + '\n\n'
                            k += 1
                        print(f"成功读取PDF文件，共{k}页，提取了前30页")
                        return content
                    # DOCX文件是ZIP格式，通常以PK开头
                    elif header.startswith(b'PK'):
                        print("检测到可能是DOCX文件，尝试读取...")
                        try:
                            doc = Document(file_path)
                            doc_text = ""
                            paragraphs = doc.paragraphs
                            for paragraph in paragraphs:
                                doc_text += paragraph.text + "\n"
                            print(f"成功读取DOCX文件，共{len(paragraphs)}段落")
                            return doc_text
                        except Exception as docx_e:
                            print(f"尝试作为DOCX读取失败：{str(docx_e)}")
            except Exception as detect_err:
                print(f"文件类型检测失败：{str(detect_err)}")
                
            print(f"不支持的文件格式：{file_path}")
            return None
    except Exception as e:
        print(f"文件读取过程中发生异常：{str(e)}")
        print(traceback.format_exc())
        return None

def split_string(s, length=15000):
    """
    将字符串按指定的长度切分成多段。

    :param s: 要切分的字符串
    :param length: 切分的字符长度，默认为15000
    :return: 切分后的字符串列表
    """
    if not s:
        print("输入字符串为空，无法切分")
        return []
    
    print(f"切分字符串，总长度：{len(s)}，单段长度：{length}")
    chunks = [s[i:i + length] for i in range(0, len(s), length)]
    print(f"切分后得到{len(chunks)}个片段")
    return chunks

def knowledge_extract(text):
    """从文本中提取知识点"""
    if not text:
        print("输入文本为空，无法提取知识点")
        return []
    
    print(f"开始提取知识点，文本长度：{len(text)}")
    try:
        # 进行文本切分
        text_list = split_string(text)
        knowledges_list = []
        for i, doc_text in enumerate(text_list):
            print(f"处理文本片段 {i+1}/{len(text_list)}")
            text_length = len(doc_text)
            knowledge_count = max(3, math.ceil(text_length / 1000))
            prompt = '''梳理抽取出上面这个文件中涉及的学习或培训知识要点；要
            求整理成一个个小的知识要点。知识点要简短精炼，
            特别注意文档中的试题、示例和案例解析内容不要做为知识点，
            每个知识点的content为当前知识点的完整描述，必须是文档的原文片段，
            可能包含多个段落，请按标准json格式返回，格式如下:
            [{"title": "车间（工段、区、队）级岗前安全培训内容",
            "content": "（一） 工作环境及危险因素； （二） 所从事工种可能遭受的职业伤害和伤亡事故； （三） 所从事工种的安全职责、操作技能及强制性标准； （四）自救互救、急救方法、疏散和现场紧急情况的处理； （五）安全设备设施、个人防护用品的使用和维护； （六）本车间（工段、区、队）安全生产状况及规章制度； （七）预防事故和职业危害的措施及应注意的安全事项； （八）有关事故案例； （九）其他需要培训的内容。"}]；
            结果不要返回除json外其它无关内容,至少包含''' + str(knowledge_count) + '''个知识要点'''

            request_text = str(doc_text) + "\n" + prompt
            messages = [
                {"role": "system",
                 "content": "你是一名学校教师或职业培训老师，你需要根据用户输入的教材或培训材料，梳理出文档材料中涉及到哪些知识点内容，这些内容是学生或企业员工需要掌握的培训要点和知识点"
                 },
                {"role": "user",
                 "content": request_text
                 },
            ]

            try:
                start_time = time.time()
                print(f"开始调用AI接口提取知识点...(片段{i+1})")
                completion = client.chat.completions.create(
                    model="qwen-plus",
                    messages=messages
                )
                result = completion.choices[0].message.content
                print(f"接口返回结果，长度：{len(result)}")

                # 提取JSON部分
                try:
                    # 尝试找到JSON数组
                    start = result.find("[")
                    end = result.rfind("]")
                    
                    if start == -1 or end == -1:
                        print(f"未找到有效的JSON数组，原始结果：{result}")
                        continue
                        
                    json_str = result[start:end + 1]
                    temp_list = json.loads(json_str)
                    print(f"成功解析JSON，获取到{len(temp_list)}个知识点")
                    knowledges_list.extend(temp_list)
                except json.JSONDecodeError as je:
                    print(f"JSON解析错误：{str(je)}")
                    print(f"原始结果：{result}")
                    continue
                    
                end_time = time.time()
                print(f"片段{i+1}处理完成，耗时：{end_time - start_time:.2f}秒")
            except Exception as api_error:
                print(f"调用AI接口错误：{str(api_error)}")
                print(traceback.format_exc())
                continue
                
        print(f"知识点提取完成，共获取{len(knowledges_list)}个知识点")
        return knowledges_list
    except Exception as e:
        print(f"第一次提取过程中出现错误：{str(e)}")
        print(traceback.format_exc())
        print("开始第二次尝试提取...")
        
        try:
            text_list = split_string(text)
            knowledges_list = []
            for i, doc_text in enumerate(text_list):
                print(f"第二次尝试：处理文本片段 {i+1}/{len(text_list)}")
                prompt = '''#注意，你是一名经验丰富的老师，请梳理抽取出上面这个文件中涉及的必考的重要知识点，整理成一个个小的知识要点。知识点要简短精炼，请按json格式返回，格式如下: [{"title": "车间（工段、区、队）级岗前安全培训内容", "content": "（一） 工作环境及危险因素； （二） 所从事工种可能遭受的职业伤害和伤亡事故； （三） 所从事工种的安全职责、操作技能及强制性标准； （四）自救互救、急救方法、疏散和现场紧急情况的处理； （五）安全设备设施、个人防护用品的使用和维护； （六）本车间（工段、区、队）安全生产状况及规章制度； （七）预防事故和职业危害的措施及应注意的安全事项； （八）有关事故案例； （九）其他需要培训的内容。"}]；结果一定不要返回除json外其它无关内容'''
                request_text = str(doc_text) + "\n" + prompt
                messages = [
                    {"role": "system",
                     "content": "你是一名学校教师或职业培训老师，你需要根据用户输入的材料，梳理出文档材料中涉及到哪些核心内容"
                     },
                    {"role": "user",
                     "content": request_text
                     }
                ]
                
                try:
                    start_time = time.time()
                    print(f"第二次尝试：调用AI接口提取知识点...(片段{i+1})")
                    completion = client.chat.completions.create(
                        model="qwen-plus",
                        messages=messages,
                        temperature=0.7  # 增加一些随机性
                    )
                    result = completion.choices[0].message.content
                    
                    # 提取JSON部分
                    try:
                        start = result.find("[")
                        end = result.rfind("]")
                        
                        if start == -1 or end == -1:
                            print(f"第二次尝试：未找到有效的JSON数组")
                            continue
                            
                        json_str = result[start:end + 1]
                        temp_list = json.loads(json_str)
                        print(f"第二次尝试：成功解析JSON，获取到{len(temp_list)}个知识点")
                        knowledges_list.extend(temp_list)
                    except json.JSONDecodeError as je:
                        print(f"第二次尝试：JSON解析错误：{str(je)}")
                        continue
                        
                    end_time = time.time()
                    print(f"第二次尝试：片段{i+1}处理完成，耗时：{end_time - start_time:.2f}秒")
                except Exception as api_error:
                    print(f"第二次尝试：调用AI接口错误：{str(api_error)}")
                    print(traceback.format_exc())
                    continue
            
            print(f"第二次尝试结束，共获取{len(knowledges_list)}个知识点")
            return knowledges_list
        except Exception as e2:
            print(f"第二次提取也失败：{str(e2)}")
            print(traceback.format_exc())
            print("知识点提取完全失败")
            return []

def generate_knowledge_graph(knowledge_point):
    """生成知识图谱"""
    import os  # 确保在函数内部导入os模块
    
    if not knowledge_point or len(knowledge_point) == 0:
        print("知识点列表为空，无法生成知识图谱")
        return "{}"
    
    print(f"开始生成知识图谱，基于{len(knowledge_point)}个知识点")
    
    # JSON 格式示例
    json_example = """{
        "name": "计算机网络核心知识点",
        "children": [
            {
                "name": "基本概念",
                "children": [
                    {
                        "name": "定义",
                        "description": "将分散的计算机等设备相互连接，以实现资源共享与信息传递。"
                    },
                    {
                        "name": "分类",
                        "children": [
                            {
                                "name": "按范围",
                                "children": [
                                    {
                                        "name": "局域网（LAN）",
                                        "description": "覆盖范围较小，通常在一个建筑物内或园区内，如学校机房网络。"
                                    },
                                    {
                                        "name": "城域网（MAN）",
                                        "description": "覆盖范围为城市区域，连接多个局域网。"
                                    },
                                    {
                                        "name": "广域网（WAN）",
                                        "description": "覆盖范围大，可跨城市、国家甚至全球，如互联网。"
                                    }
                                ]
                            },
                            {
                                "name": "按拓扑结构",
                                "children": [
                                    {
                                        "name": "总线型",
                                        "description": "所有设备连接在一条总线上，如早期的同轴电缆网络。"
                                    },
                                    {
                                        "name": "星型",
                                        "description": "以中心节点为核心，其他设备通过线缆连接到中心节点，目前应用广泛。"
                                    },
                                    {
                                        "name": "环型",
                                        "description": "设备连接成环，数据沿环单向传输。"
                                    }
                                ]
                            }
                        ]
                    }
                ]
            },
            {
                "name": "网络协议",
                "children": [
                    {
                        "name": "网络层",
                        "children": [
                            {
                                "name": "IP协议",
                                "description": "负责寻址和路由，为数据包分配 IP 地址并选择传输路径。"
                            }
                        ]
                    },
                    {
                        "name": "传输层",
                        "children": [
                            {
                                "name": "TCP",
                                "description": "保证可靠传输，通过三次握手建立连接，确保数据按序、无差错到达。"
                            },
                            {
                                "name": "UDP",
                                "description": "用于快速传输，不保证数据可靠性，适用于实时性要求高但对数据准确性要求稍低的场景，如视频流。"
                            }
                        ]
                    },
                    {
                        "name": "应用层",
                        "children": [
                            {
                                "name": "HTTP",
                                "description": "用于网页传输，浏览器与服务器之间通信的协议。"
                            },
                            {
                                "name": "SMTP",
                                "description": "邮件发送协议，负责将邮件从客户端发送到邮件服务器。"
                            }
                        ]
                    }
                ]
            },
            {
                "name": "网络管理",
                "children": [
                    {
                        "name": "性能管理",
                        "description": "监测网络性能，如带宽利用率、延迟等，确保网络高效运行。"
                    },
                    {
                        "name": "故障管理",
                        "description": "检测并排除故障，通过监控设备状态和日志，快速定位和解决问题。"
                    },
                    {
                        "name": "配置管理",
                        "description": "管理设备配置，包括设备参数设置、软件升级等，保证网络设备正常工作。"
                    }
                ]
            }
        ]
    }"""

    try:
        # 系统提示
        prompt_system = "你是一名 json 格式数据处理专家，能够理解分析处理各种复杂 json 数据，并将其按照要求重新整合成符合用户需要的数据"

        # 用户提示
        try:
            prompt_1 = "现有一批数据为知识点抽取结果：" + json.dumps(knowledge_point, ensure_ascii=False) + \
                      "将以上知识点抽取结果数据转换为知识图谱数据，返回的知识图谱数据需要对原文本进行知识点逻辑分层，返回结果仍为标准 json 格式，其格式示例为："
            prompt_tips = "另外请注意，返回的 json 数据中不需要对父知识点设置 description 字段。"
        except Exception as json_err:
            print(f"JSON序列化错误：{str(json_err)}")
            print(traceback.format_exc())
            # 尝试简化知识点对象
            simplified_points = []
            for point in knowledge_point:
                try:
                    simple_point = {
                        "title": str(point.get("title", "未知标题"))[:100],
                        "content": str(point.get("content", "未知内容"))[:200]
                    }
                    simplified_points.append(simple_point)
                except Exception:
                    continue
            
            prompt_1 = "现有一批数据为知识点抽取结果：" + json.dumps(simplified_points, ensure_ascii=False) + \
                      "将以上知识点抽取结果数据转换为知识图谱数据，返回的知识图谱数据需要对原文本进行知识点逻辑分层，返回结果仍为标准 json 格式，其格式示例为："
            prompt_tips = "另外请注意，返回的 json 数据中不需要对父知识点设置 description 字段。"

        # 调用 OpenAI API
        print("开始调用AI接口生成知识图谱...")
        try:
            completion = client.chat.completions.create(
                model="qwen-plus",
                messages=[
                    {"role": "system", "content": prompt_system},
                    {"role": "user", "content": prompt_1 + json_example + prompt_tips},
                ],
                temperature=0.5
            )
            
            result = completion.choices[0].message.content
            print("成功获取到知识图谱数据")
            
            # 确保返回的是有效的JSON
            try:
                # 找到JSON对象
                start = result.find("{")
                end = result.rfind("}")
                
                if start == -1 or end == -1:
                    print("未找到有效的JSON对象，返回简单对象")
                    return json.dumps({"name": "知识图谱", "children": []})
                    
                json_str = result[start:end + 1]
                # 验证JSON是否有效
                json.loads(json_str)
                print("验证知识图谱JSON格式正确")
                
                # 保存知识图谱数据到文件
                try:
                    # 使用绝对路径保存文件或者可以完全跳过保存步骤
                    # output_file = "knowledge_graph.json"
                    # with open(output_file, 'w', encoding='utf-8') as file:
                    #     file.write(json_str)
                    # print(f"知识图谱数据已保存到 {output_file}")
                    print("知识图谱数据生成成功，跳过保存到文件")
                except Exception as file_err:
                    print(f"保存文件失败：{str(file_err)}")
                
                return json_str
            except json.JSONDecodeError as je:
                print(f"生成的知识图谱JSON无效：{str(je)}")
                print(f"原始结果：{result}")
                # 返回一个简单的有效JSON
                return json.dumps({"name": "知识图谱", "children": []})
                
        except Exception as api_err:
            print(f"调用AI接口生成知识图谱失败：{str(api_err)}")
            print(traceback.format_exc())
            return json.dumps({"name": "知识图谱", "children": []})
            
    except Exception as e:
        print(f"知识图谱生成过程中发生异常：{str(e)}")
        print(traceback.format_exc())
        return json.dumps({"name": "知识图谱", "children": []})

if __name__ == "__main__":
    file_path = "./uploads/course_files/Cae6d63_090.pdf"
    text = read_file_content(file_path)
    if text:
        knowledges = knowledge_extract(text)
        print("抽取的知识点：")
        print(json.dumps(knowledges, ensure_ascii=False, indent=4))

        print("知识图谱：")
        print(generate_knowledge_graph(knowledges))